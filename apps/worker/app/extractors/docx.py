"""
DOCX extractor — uses python-docx for Word document text extraction.
Extracts paragraphs and tables with section headers preserved.
"""

import logging

from docx import Document
from docx.table import Table

from . import ExtractionResult

logger = logging.getLogger(__name__)


def _extract_table_text(table: Table) -> str:
    """Convert a DOCX table to readable text format."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


async def extract_docx(file_path: str) -> ExtractionResult:
    """Extract text from a DOCX file, preserving headings and tables."""
    doc = Document(file_path)
    sections: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]  # Strip namespace

        if tag == "p":
            # Paragraph — check if it's a heading
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()
                    if not text:
                        break
                    if para.style and para.style.name.startswith("Heading"):
                        level = para.style.name.replace("Heading ", "").strip()
                        prefix = "#" * int(level) if level.isdigit() else "#"
                        sections.append(f"{prefix} {text}")
                    else:
                        sections.append(text)
                    break

        elif tag == "tbl":
            # Table
            for table in doc.tables:
                if table._element is element:
                    table_text = _extract_table_text(table)
                    if table_text.strip():
                        sections.append(f"\n[TABLE]\n{table_text}\n[/TABLE]")
                    break

    full_text = "\n\n".join(sections)
    word_count = len(full_text.split())

    # Estimate page count (~500 words per page for Word docs)
    page_count = max(1, word_count // 500)

    logger.info(f"DOCX extracted: ~{page_count} pages, {word_count} words — {file_path}")

    return ExtractionResult(
        text=full_text,
        page_count=page_count,
        word_count=word_count,
        metadata={"extractor": "python-docx"},
    )
