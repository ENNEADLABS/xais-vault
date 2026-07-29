"""
DOCX Sections — element dataclasses, Markdown parser, and inline formatters.

Extracted from docx_builder.py for the 200-line-per-file rule.
Contains the element model and rendering helpers used by build_docx().
"""

import re
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Pt

from .docx_styles import (
    FONT_NAME,
    HEADING_COLOR_1,
    HEADING_COLOR_2,
)

# ─── Element dataclasses ─────────────────────────────────────────


@dataclass
class Heading:
    level: int
    text: str


@dataclass
class ListItem:
    text: str
    ordered: bool = False


@dataclass
class TableElement:
    rows: list[list[str]] = field(default_factory=list)


@dataclass
class Quote:
    text: str


@dataclass
class HorizontalRule:
    pass


@dataclass
class TextParagraph:
    text: str


Element = Heading | ListItem | TableElement | Quote | HorizontalRule | TextParagraph


# ─── Markdown parser ──────────────────────────────────────────────


def _parse_table(lines: list[str]) -> list[list[str]]:
    """Parse Markdown table lines into a 2D list, excluding separator rows."""
    rows = []
    for line in lines:
        if re.match(r"^\|[-| :]+\|$", line.strip()):
            continue  # Skip separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def _parse_markdown_lines(content: str) -> list[Element]:
    """Parse Markdown content into a list of typed elements."""
    elements: list[Element] = []
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("#### "):
            elements.append(Heading(level=3, text=line[5:]))
        elif line.startswith("### "):
            elements.append(Heading(level=3, text=line[4:]))
        elif line.startswith("## "):
            elements.append(Heading(level=2, text=line[3:]))
        elif line.startswith("# "):
            elements.append(Heading(level=1, text=line[2:]))
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = _parse_table(table_lines)
            if rows:
                elements.append(TableElement(rows=rows))
            continue
        elif line.startswith("- ") or line.startswith("* "):
            elements.append(ListItem(text=line[2:], ordered=False))
        elif re.match(r"^\d+\. ", line):
            elements.append(ListItem(text=re.sub(r"^\d+\. ", "", line), ordered=True))
        elif line.startswith("> "):
            elements.append(Quote(text=line[2:]))
        elif line.strip() == "---":
            elements.append(HorizontalRule())
        elif line.strip():
            elements.append(TextParagraph(text=line))
        i += 1

    return elements


# ─── Inline formatting ────────────────────────────────────────────


def _add_formatted_runs(paragraph, text: str) -> None:
    """Apply **bold** and *italic* Markdown inline formatting as DOCX runs."""
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*)"
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part:
            paragraph.add_run(part)


# ─── Document element renderers ──────────────────────────────────


def _add_heading_paragraph(doc: Document, element: Heading) -> None:
    p = doc.add_paragraph(style=f"Heading {min(element.level, 3)}")
    run = p.add_run(element.text)
    run.font.name = FONT_NAME
    if element.level == 1:
        run.font.size = Pt(24)
        run.font.color.rgb = HEADING_COLOR_1
    elif element.level == 2:
        run.font.size = Pt(18)
        run.font.color.rgb = HEADING_COLOR_2
    else:
        run.font.size = Pt(14)
        run.font.color.rgb = HEADING_COLOR_2


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    """Add a styled table with a bold header row."""
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= col_count:
                break
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            if row_idx == 0:
                run.bold = True

    doc.add_paragraph()  # spacing after table
