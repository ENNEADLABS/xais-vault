"""
DOCX Builder — converts Claude's Markdown output to a professional Word document.

Handles: headings, lists, tables, bold/italic, blockquotes, cover page, TOC.
Uses python-docx. No external markdown library needed — custom line parser.
"""

import io
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .docx_sections import (
    Element,
    Heading,
    HorizontalRule,
    ListItem,
    Quote,
    TableElement,
    TextParagraph,
    _add_formatted_runs,
    _add_heading_paragraph,
    _add_table,
    _parse_markdown_lines,
)
from .docx_styles import (
    FONT_NAME,
    HEADING_COLOR_1,
    HEADING_COLOR_2,
    MUTED_COLOR,
    TYPE_LABELS,
)


def _add_toc(doc: Document) -> None:
    """Insert a Word TOC field (auto-populated when document is opened in Word)."""
    heading = doc.add_paragraph("Table des matières")
    heading.style = "Heading 2"

    para = doc.add_paragraph()
    run = para.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instr)

    run3 = para.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fld_end)


def _add_cover_page(
    doc: Document,
    deliverable_type: str,
    workspace_name: str,
    target_company: str | None,
) -> None:
    """Add a professional cover page."""
    type_label = TYPE_LABELS.get(deliverable_type, deliverable_type.upper().replace("_", " "))

    for _ in range(8):
        doc.add_paragraph()

    title_p = doc.add_paragraph("XAIS Vault")
    title_p.alignment = 1  # CENTER
    run = title_p.runs[0] if title_p.runs else title_p.add_run("XAIS Vault")
    run.font.name = FONT_NAME
    run.font.size = Pt(28)
    run.font.color.rgb = HEADING_COLOR_1
    run.bold = True

    type_p = doc.add_paragraph(type_label)
    type_p.alignment = 1
    type_run = type_p.runs[0] if type_p.runs else type_p.add_run(type_label)
    type_run.font.name = FONT_NAME
    type_run.font.size = Pt(20)
    type_run.font.color.rgb = HEADING_COLOR_2
    type_run.bold = True

    doc.add_paragraph()

    name_p = doc.add_paragraph(workspace_name)
    name_p.alignment = 1
    name_run = name_p.runs[0] if name_p.runs else name_p.add_run(workspace_name)
    name_run.font.name = FONT_NAME
    name_run.font.size = Pt(16)
    name_run.bold = True

    if target_company and target_company != workspace_name:
        co_p = doc.add_paragraph(target_company)
        co_p.alignment = 1
        co_run = co_p.runs[0] if co_p.runs else co_p.add_run(target_company)
        co_run.font.name = FONT_NAME
        co_run.font.size = Pt(14)

    for _ in range(6):
        doc.add_paragraph()

    date_p = doc.add_paragraph(f"Date : {datetime.now().strftime('%d %B %Y')}")
    date_p.alignment = 1

    conf_p = doc.add_paragraph("Généré par XAIS Vault — CONFIDENTIEL")
    conf_p.alignment = 1
    conf_run = conf_p.runs[0] if conf_p.runs else conf_p.add_run()
    conf_run.font.size = Pt(9)
    conf_run.font.color.rgb = MUTED_COLOR


# ─── Main entry point ─────────────────────────────────────────────


def build_docx(
    markdown_content: str,
    deliverable_type: str,
    workspace_name: str,
    target_company: str | None = None,
) -> bytes:
    """Convert Markdown content to a professional DOCX. Returns the file bytes."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)

    _add_cover_page(doc, deliverable_type, workspace_name, target_company)
    doc.add_page_break()

    if deliverable_type in ("investment_memo", "dd_report"):
        _add_toc(doc)
        doc.add_page_break()

    elements = _parse_markdown_lines(markdown_content)
    for element in elements:
        if isinstance(element, Heading):
            _add_heading_paragraph(doc, element)
        elif isinstance(element, ListItem):
            list_style = "List Bullet" if not element.ordered else "List Number"
            p = doc.add_paragraph(style=list_style)
            _add_formatted_runs(p, element.text)
        elif isinstance(element, TableElement):
            _add_table(doc, element.rows)
        elif isinstance(element, Quote):
            p = doc.add_paragraph(style="Quote")
            run = p.add_run(element.text)
            run.italic = True
        elif isinstance(element, HorizontalRule):
            doc.add_paragraph()
        elif isinstance(element, TextParagraph):
            p = doc.add_paragraph()
            _add_formatted_runs(p, element.text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# Re-exports for backward compatibility (tests import from docx_builder)
__all__ = [
    "Element",
    "Heading",
    "HorizontalRule",
    "ListItem",
    "Quote",
    "TableElement",
    "TextParagraph",
    "_add_formatted_runs",
    "_parse_markdown_lines",
    "build_docx",
]
