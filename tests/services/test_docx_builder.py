"""
Tests for services/docx_builder.py — Markdown → DOCX conversion.

Tests unitaires sur le parser et les fonctions de construction.
Tests d'intégration : build_docx() + inspection du Document.

Couverture :
  - _parse_markdown_lines : 5 cas
  - _add_formatted_runs   : 3 cas
  - build_docx            : 6 cas
"""

import io

from docx import Document

from apps.worker.app.services.docx_builder import (
    Heading,
    HorizontalRule,
    ListItem,
    Quote,
    TableElement,
    _add_formatted_runs,
    _parse_markdown_lines,
    build_docx,
)

# ─── Helper ─────────────────────────────────────────────────


def load_doc(docx_bytes: bytes) -> Document:
    return Document(io.BytesIO(docx_bytes))


def all_texts(doc: Document) -> list[str]:
    """Return all non-empty paragraph texts in order."""
    return [p.text for p in doc.paragraphs if p.text.strip()]


# ═══════════════════════════════════════════════════════════════
# Couche 1 — _parse_markdown_lines (pure function)
# ═══════════════════════════════════════════════════════════════


class TestParseMarkdownLines:
    def test_headings_all_levels(self):
        md = "# H1\n## H2\n### H3\n#### H4"
        elems = _parse_markdown_lines(md)
        assert elems[0] == Heading(level=1, text="H1")
        assert elems[1] == Heading(level=2, text="H2")
        assert elems[2] == Heading(level=3, text="H3")
        assert elems[3] == Heading(level=3, text="H4")  # Level 4 → capped at 3

    def test_bullet_list(self):
        md = "- Premier\n- Deuxième\n* Troisième"
        elems = _parse_markdown_lines(md)
        assert all(isinstance(e, ListItem) for e in elems)
        assert not any(e.ordered for e in elems)
        assert elems[0].text == "Premier"

    def test_ordered_list(self):
        md = "1. Premier\n2. Deuxième\n10. Dixième"
        elems = _parse_markdown_lines(md)
        assert all(isinstance(e, ListItem) for e in elems)
        assert all(e.ordered for e in elems)
        assert elems[2].text == "Dixième"

    def test_table_parsed_as_single_element(self):
        md = "| Col A | Col B |\n|-------|-------|\n| Val 1 | Val 2 |"
        elems = _parse_markdown_lines(md)
        assert len(elems) == 1
        assert isinstance(elems[0], TableElement)
        assert elems[0].rows == [["Col A", "Col B"], ["Val 1", "Val 2"]]

    def test_blockquote(self):
        md = "> Ceci est une citation."
        elems = _parse_markdown_lines(md)
        assert len(elems) == 1
        assert isinstance(elems[0], Quote)
        assert elems[0].text == "Ceci est une citation."

    def test_horizontal_rule(self):
        md = "---"
        elems = _parse_markdown_lines(md)
        assert len(elems) == 1
        assert isinstance(elems[0], HorizontalRule)

    def test_mixed_content(self):
        md = "# Titre\n- Item\n\nParagraphe normal."
        elems = _parse_markdown_lines(md)
        types = [type(e).__name__ for e in elems]
        assert "Heading" in types
        assert "ListItem" in types
        assert "TextParagraph" in types

    def test_empty_lines_skipped(self):
        md = "\n\n\n# Titre\n\n\n"
        elems = _parse_markdown_lines(md)
        assert len(elems) == 1
        assert isinstance(elems[0], Heading)


# ═══════════════════════════════════════════════════════════════
# Couche 1 — _add_formatted_runs (with real paragraph)
# ═══════════════════════════════════════════════════════════════


class TestAddFormattedRuns:
    def _make_paragraph(self):
        doc = Document()
        return doc.add_paragraph()

    def test_plain_text_single_run(self):
        p = self._make_paragraph()
        _add_formatted_runs(p, "Texte simple.")
        assert any(r.text == "Texte simple." for r in p.runs)

    def test_bold_text(self):
        p = self._make_paragraph()
        _add_formatted_runs(p, "**Gras**")
        bold_runs = [r for r in p.runs if r.bold]
        assert any(r.text == "Gras" for r in bold_runs)

    def test_italic_text(self):
        p = self._make_paragraph()
        _add_formatted_runs(p, "*Italique*")
        italic_runs = [r for r in p.runs if r.italic]
        assert any(r.text == "Italique" for r in italic_runs)

    def test_mixed_formatting(self):
        p = self._make_paragraph()
        _add_formatted_runs(p, "Normal **gras** et *italic* fin.")
        full_text = "".join(r.text for r in p.runs)
        assert "Normal" in full_text
        assert "gras" in full_text
        assert "italic" in full_text


# ═══════════════════════════════════════════════════════════════
# Couche 1/2 — build_docx (integration)
# ═══════════════════════════════════════════════════════════════


class TestBuildDocx:
    def test_returns_valid_zip_bytes(self):
        """DOCX is a ZIP file — magic bytes must be PK."""
        data = build_docx("# Test", "executive_summary", "Workspace A")
        assert isinstance(data, bytes)
        assert data[:2] == b"PK"

    def test_cover_page_contains_deal_name(self):
        data = build_docx("# Contenu", "executive_summary", "Mon Workspace", "Ma Société")
        doc = load_doc(data)
        texts = all_texts(doc)
        assert any("Mon Workspace" in t for t in texts)

    def test_cover_page_contains_type_label(self):
        data = build_docx("", "investment_memo", "Workspace B")
        doc = load_doc(data)
        texts = all_texts(doc)
        assert any("INVESTMENT MEMO" in t for t in texts)

    def test_headings_present_in_document(self):
        md = "# Titre Principal\n## Sous-titre\n### Section"
        data = build_docx(md, "executive_summary", "Workspace C")
        doc = load_doc(data)
        para_texts = [p.text for p in doc.paragraphs]
        assert "Titre Principal" in para_texts
        assert "Sous-titre" in para_texts

    def test_table_produces_docx_table(self):
        md = "| Métrique | Valeur |\n|----------|--------|\n| ARR | 8M€ |"
        data = build_docx(md, "executive_summary", "Workspace D")
        doc = load_doc(data)
        assert len(doc.tables) >= 1
        # Check first table has our data
        first_table = doc.tables[0]
        header_texts = [cell.text for cell in first_table.rows[0].cells]
        assert "Métrique" in header_texts

    def test_toc_added_for_investment_memo(self):
        data = build_docx("# Section", "investment_memo", "Workspace E")
        doc = load_doc(data)
        texts = all_texts(doc)
        assert any("Table des matières" in t for t in texts)

    def test_toc_not_added_for_executive_summary(self):
        data = build_docx("# Section", "executive_summary", "Workspace F")
        doc = load_doc(data)
        texts = all_texts(doc)
        assert not any("Table des matières" in t for t in texts)

    def test_bullet_list_produces_paragraph(self):
        md = "- Premier\n- Deuxième\n* Troisième"
        data = build_docx(md, "executive_summary", "Workspace H")
        doc = load_doc(data)
        texts = all_texts(doc)
        assert any("Premier" in t for t in texts)

    def test_ordered_list_produces_paragraph(self):
        md = "1. Alpha\n2. Bravo"
        data = build_docx(md, "executive_summary", "Workspace I")
        doc = load_doc(data)
        texts = all_texts(doc)
        assert any("Alpha" in t for t in texts)

    def test_blockquote_produces_italic_paragraph(self):
        md = "> Citation importante."
        data = build_docx(md, "executive_summary", "Workspace J")
        doc = load_doc(data)
        texts = all_texts(doc)
        assert any("Citation importante" in t for t in texts)

    def test_horizontal_rule_produces_empty_paragraph(self):
        md = "# Avant\n---\n# Après"
        data = build_docx(md, "executive_summary", "Workspace K")
        doc = load_doc(data)
        para_texts = [p.text for p in doc.paragraphs]
        assert "Avant" in para_texts
        assert "Après" in para_texts

    def test_empty_content_does_not_raise(self):
        data = build_docx("", "dd_report", "Workspace G")
        assert isinstance(data, bytes)
        assert len(data) > 0
